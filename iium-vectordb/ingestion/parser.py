"""
Parse raw documents (PDF, DOCX, TXT) into clean plain text.

Each file produces a ParsedDocument with:
  - source: original file path (str)
  - pages: list of (page_number, text) tuples
  - full_text: concatenated text of all pages
"""
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    source: str
    pages: list[tuple[int, str]] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(text for _, text in self.pages)

    @property
    def page_count(self) -> int:
        return len(self.pages)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> ParsedDocument:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Install pdfplumber: pip install pdfplumber")

    doc = ParsedDocument(source=str(path))
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            text = _clean_text(text)
            if text.strip():
                doc.pages.append((i, text))
    return doc


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")

    document = Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # Treat the whole DOCX as a single "page"
    text = _clean_text("\n".join(paragraphs))
    doc = ParsedDocument(source=str(path))
    if text.strip():
        doc.pages.append((1, text))
    return doc


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------

def _parse_txt(path: Path) -> ParsedDocument:
    text = _clean_text(path.read_text(encoding="utf-8", errors="replace"))
    doc = ParsedDocument(source=str(path))
    if text.strip():
        doc.pages.append((1, text))
    return doc


# ---------------------------------------------------------------------------
# Text cleaning
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    # Collapse multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove non-printable characters except newlines/tabs
    text = re.sub(r"[^\x09\x0A\x20-\x7E\u00A0-\uFFFF]", " ", text)
    # Collapse multiple spaces
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

PARSERS = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".doc": _parse_docx,
    ".txt": _parse_txt,
}


def parse_file(path: Path) -> Optional[ParsedDocument]:
    ext = path.suffix.lower()
    parser = PARSERS.get(ext)
    if parser is None:
        logger.warning("No parser for extension %s: %s", ext, path.name)
        return None
    try:
        doc = parser(path)
        logger.info("Parsed %s (%d pages, %d chars)", path.name, doc.page_count, len(doc.full_text))
        return doc
    except Exception as exc:
        logger.error("Failed to parse %s: %s", path.name, exc)
        return None


def parse_directory(raw_dir: Path) -> list[ParsedDocument]:
    files = [p for p in raw_dir.rglob("*") if p.is_file() and p.suffix.lower() in PARSERS]
    logger.info("Found %d parseable files in %s", len(files), raw_dir)
    docs = []
    for f in files:
        doc = parse_file(f)
        if doc and doc.full_text.strip():
            docs.append(doc)
    logger.info("Successfully parsed %d / %d files", len(docs), len(files))
    return docs


if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
    raw_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("ingestion/raw_docs")
    docs = parse_directory(raw_dir)
    for d in docs:
        print(f"  {Path(d.source).name}: {d.page_count} pages, {len(d.full_text)} chars")
