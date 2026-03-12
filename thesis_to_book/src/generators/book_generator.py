"""Generate a formatted Word (.docx) academic book from rewritten chapters."""

from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.util import Twips

from ..processors.ai_rewriter import RewrittenChapter


# ---------------------------------------------------------------------------
# XML helpers for Word-native TOC field
# ---------------------------------------------------------------------------

def _add_toc_field(doc: Document) -> None:
    """Insert a Word TOC field that updates when the document is opened."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_end)


def _add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ---------------------------------------------------------------------------
# BookGenerator
# ---------------------------------------------------------------------------

class BookGenerator:
    """
    Compile a list of RewrittenChapters into a single formatted Word document.

    The generated book contains:
      - Title page
      - Table of Contents (Word native field)
      - Editors' Preface
      - Chapters (abstract, introduction, body, conclusion)
      - Consolidated bibliography
      - Index
    """

    def __init__(self, book_config: dict[str, Any]):
        self._cfg = book_config

    def generate(
        self,
        chapters: list[RewrittenChapter],
        output_path: str,
        include_toc: bool = True,
        include_bibliography: bool = True,
        include_index: bool = True,
    ) -> str:
        doc = Document()
        self._set_page_margins(doc)
        self._set_default_styles(doc)

        self._add_title_page(doc)
        _add_page_break(doc)

        if include_toc:
            doc.add_heading("Table of Contents", level=1)
            _add_toc_field(doc)
            _add_page_break(doc)

        self._add_preface(doc, chapters)
        _add_page_break(doc)

        all_refs: list[str] = []
        index_map: dict[str, list[int]] = defaultdict(list)

        for chapter in chapters:
            self._add_chapter(doc, chapter)
            all_refs.extend(chapter.references)
            for term in chapter.index_terms:
                index_map[term].append(chapter.chapter_number)
            _add_page_break(doc)

        if include_bibliography:
            self._add_bibliography(doc, all_refs)
            _add_page_break(doc)

        if include_index and index_map:
            self._add_index(doc, index_map)

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return str(out)

    # ------------------------------------------------------------------
    # Page / style setup
    # ------------------------------------------------------------------

    def _set_page_margins(self, doc: Document) -> None:
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

    def _set_default_styles(self, doc: Document) -> None:
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = Pt(18)

    # ------------------------------------------------------------------
    # Title page
    # ------------------------------------------------------------------

    def _add_title_page(self, doc: Document) -> None:
        cfg = self._cfg
        doc.add_paragraph()
        doc.add_paragraph()

        title_para = doc.add_paragraph(cfg.get("title", "Academic Book"))
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.runs[0]
        run.bold = True
        run.font.size = Pt(24)

        subtitle = cfg.get("subtitle", "")
        if subtitle:
            sub_para = doc.add_paragraph(subtitle)
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.runs[0].font.size = Pt(16)
            sub_para.runs[0].italic = True

        doc.add_paragraph()

        editors = cfg.get("editors", [])
        if editors:
            label = "Editors:" if len(editors) > 1 else "Editor:"
            ed_para = doc.add_paragraph(label)
            ed_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for ed in editors:
                p = doc.add_paragraph(ed)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        inst = cfg.get("institution", "")
        if inst:
            ip = doc.add_paragraph(inst)
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.runs[0].font.size = Pt(11)

        year = cfg.get("year", "")
        if year:
            yp = doc.add_paragraph(str(year))
            yp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------------------------
    # Preface
    # ------------------------------------------------------------------

    def _add_preface(self, doc: Document, chapters: list[RewrittenChapter]) -> None:
        doc.add_heading("Editors' Preface", level=1)
        subject = self._cfg.get("subject_area", "the subject area")
        editors = self._cfg.get("editors", ["The Editors"])
        editor_str = " and ".join(editors)

        doc.add_paragraph(
            f"This volume brings together original doctoral research in the field of "
            f"{subject}. Each chapter has been adapted from a doctoral thesis and "
            f"revised to form a coherent contribution to a unified academic work."
        )
        contrib_list = ", ".join(
            f"\"{ch.chapter_title}\" by {ch.author}" for ch in chapters
        )
        doc.add_paragraph(
            f"The chapters collected here are: {contrib_list}. "
            "Together they represent a comprehensive survey of current research "
            f"in {subject}."
        )
        doc.add_paragraph(f"— {editor_str}")

    # ------------------------------------------------------------------
    # Chapter
    # ------------------------------------------------------------------

    def _add_chapter(self, doc: Document, chapter: RewrittenChapter) -> None:
        # Chapter heading
        heading_text = f"Chapter {chapter.chapter_number}: {chapter.chapter_title}"
        doc.add_heading(heading_text, level=1)

        author_para = doc.add_paragraph(f"Author: {chapter.author}")
        author_para.runs[0].italic = True
        author_para.runs[0].font.size = Pt(11)

        # Abstract box
        if chapter.abstract:
            doc.add_heading("Abstract", level=3)
            abstract_para = doc.add_paragraph(chapter.abstract)
            abstract_para.paragraph_format.left_indent = Inches(0.4)
            abstract_para.paragraph_format.right_indent = Inches(0.4)
            abstract_para.runs[0].font.size = Pt(11)
            abstract_para.runs[0].italic = True

        # Introduction
        if chapter.introduction:
            doc.add_heading("Introduction", level=2)
            self._add_body_text(doc, chapter.introduction)

        # Body sections
        for heading, content in chapter.body_sections:
            clean_heading = self._clean_heading(heading)
            doc.add_heading(clean_heading, level=2)
            self._add_body_text(doc, content)

        # Conclusion
        if chapter.conclusion:
            doc.add_heading("Conclusion", level=2)
            self._add_body_text(doc, chapter.conclusion)

    def _add_body_text(self, doc: Document, text: str) -> None:
        for para_text in text.split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)

    def _clean_heading(self, heading: str) -> str:
        """Remove leading numbering like '2.1 ' from headings."""
        return re.sub(r"^\d+(\.\d+)*\s+", "", heading).strip()

    # ------------------------------------------------------------------
    # Bibliography
    # ------------------------------------------------------------------

    def _add_bibliography(self, doc: Document, refs: list[str]) -> None:
        doc.add_heading("Bibliography", level=1)
        seen: set[str] = set()
        unique_refs = [r for r in refs if r and r not in seen and not seen.add(r)]  # type: ignore[func-returns-value]
        unique_refs.sort()
        for ref in unique_refs:
            para = doc.add_paragraph(ref, style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.5)

    # ------------------------------------------------------------------
    # Index
    # ------------------------------------------------------------------

    def _add_index(self, doc: Document, index_map: dict[str, list[int]]) -> None:
        doc.add_heading("Index", level=1)
        sorted_terms = sorted(index_map.keys(), key=lambda t: t.lower())

        current_letter = ""
        for term in sorted_terms:
            first_letter = term[0].upper()
            if first_letter != current_letter:
                current_letter = first_letter
                letter_para = doc.add_paragraph(current_letter)
                letter_para.runs[0].bold = True
                letter_para.runs[0].font.size = Pt(13)

            chapters_str = ", ".join(f"Ch. {n}" for n in sorted(index_map[term]))
            entry_para = doc.add_paragraph(f"{term}  —  {chapters_str}")
            entry_para.paragraph_format.left_indent = Inches(0.3)
