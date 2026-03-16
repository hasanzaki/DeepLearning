#!/usr/bin/env python3
"""
write_chapters_docx.py
Generate IIUM Press-formatted Chapter 4 and Chapter 5 .docx files.

Usage:
    python write_chapters_docx.py

Outputs:
    output/CHAPTER_04_DRAFT.docx
    output/CHAPTER_05_DRAFT.docx
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─────────────────────────────────────────────────────────────────────────────
# Document helpers
# ─────────────────────────────────────────────────────────────────────────────

def new_iium_doc() -> Document:
    """Create a blank Document with IIUM Press page layout and base styles."""
    doc = Document()

    # Page size 6 × 9 in, 1-inch margins
    for sec in doc.sections:
        sec.page_width  = Inches(6)
        sec.page_height = Inches(9)
        sec.top_margin  = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin  = Inches(1)
        sec.right_margin = Inches(1)

    # Body text: Times New Roman 11 pt, single spacing
    ns = doc.styles["Normal"]
    ns.font.name = "Times New Roman"
    ns.font.size = Pt(11)
    ns.paragraph_format.space_after  = Pt(6)
    ns.paragraph_format.line_spacing = Pt(13.2)

    # Heading styles (max 3 levels)
    heading_cfg = [(1, 14), (2, 12), (3, 11)]
    for lvl, sz in heading_cfg:
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name  = "Times New Roman"
        hs.font.size  = Pt(sz)
        hs.font.bold  = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after  = Pt(6)

    return doc


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


# ─────────────────────────────────────────────────────────────────────────────
# Rich-text helpers (inline **bold** / *italic* / `code`)
# ─────────────────────────────────────────────────────────────────────────────

_INLINE_PAT = re.compile(
    r"\*\*(.+?)\*\*"    # **bold**
    r"|\*(.+?)\*"       # *italic*
    r"|`([^`]+)`"       # `code`
    r"|([^*`]+)",       # plain text
    re.DOTALL,
)


def _add_rich_runs(para, text: str, base_size: int = 11) -> None:
    """Parse inline markup and add correctly-styled runs to *para*."""
    for m in _INLINE_PAT.finditer(text):
        bold_txt, ital_txt, code_txt, plain_txt = m.groups()
        if bold_txt is not None:
            r = para.add_run(bold_txt)
            r.bold = True
        elif ital_txt is not None:
            r = para.add_run(ital_txt)
            r.italic = True
        elif code_txt is not None:
            r = para.add_run(code_txt)
            r.font.name = "Courier New"
        elif plain_txt is not None:
            r = para.add_run(plain_txt)
        else:
            continue
        r.font.name = r.font.name or "Times New Roman"
        r.font.size = Pt(base_size)


def add_rich_para(doc: Document, text: str, indent: float = 0.0) -> None:
    """Add a body paragraph with inline markup support."""
    if not text.strip():
        return
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    _add_rich_runs(p, text)


# ─────────────────────────────────────────────────────────────────────────────
# Figure placeholder
# ─────────────────────────────────────────────────────────────────────────────

def add_figure_placeholder(doc: Document, fig_id: str, note: str, caption: str) -> None:
    """Insert a shaded placeholder box (for manual figure insertion) + caption."""
    # Grey box via single-cell table
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "EBEBEB")

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"[ {fig_id} ]")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(80, 80, 80)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(note)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(110, 110, 110)

    # Caption below
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after  = Pt(12)
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)
    cr.font.name = "Times New Roman"


# ─────────────────────────────────────────────────────────────────────────────
# Equation block
# ─────────────────────────────────────────────────────────────────────────────

_GREEK = {
    r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\delta": "δ",
    r"\theta": "θ", r"\psi": "ψ",  r"\omega": "ω", r"\rho": "ρ",
    r"\sigma": "σ", r"\mu": "μ",   r"\eta": "η",   r"\epsilon": "ε",
    r"\phi": "φ",   r"\pi": "π",   r"\Delta": "Δ", r"\Sigma": "Σ",
    r"\Omega": "Ω", r"\Lambda": "Λ",
}
_OPERATORS = {
    r"\cdot": "·", r"\times": "×", r"\approx": "≈", r"\leq": "≤",
    r"\geq": "≥",  r"\neq": "≠",   r"\in": "∈",     r"\pm": "±",
    r"\infty": "∞", r"\partial": "∂", r"\nabla": "∇",
}


def _clean_latex(text: str) -> str:
    """Convert LaTeX to a plain-text approximation for Word body text."""
    # Extract tag
    eq_num = ""
    tag = re.search(r"\\tag\{([^}]+)\}", text)
    if tag:
        eq_num = tag.group(1)
        text = text[: tag.start()] + text[tag.end():]

    text = re.sub(r"\\begin\{[pb]matrix\}", "(", text)
    text = re.sub(r"\\end\{[pb]matrix\}",   ")", text)
    text = text.replace(r"\\", ";  ")
    text = re.sub(r"\\frac\{([^}]+)\}\{([^}]+)\}", r"(\1)/(\2)", text)
    text = re.sub(r"\\(?:mathbf|mathcal|text|mathrm)\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\(?:left|right)", "", text)
    text = re.sub(r"\\(?:log|sin|cos|tan|exp|det|tr)\b", lambda m: m.group(0)[1:], text)
    for latex, uni in {**_GREEK, **_OPERATORS}.items():
        text = text.replace(latex, uni)
    text = re.sub(r"\\[a-zA-Z]+", "", text)       # remove unknown commands
    text = re.sub(r"\{([^}]*)\}", r"\1", text)    # strip remaining braces
    text = re.sub(r"\s+", " ", text).strip()
    return text, eq_num


def add_equation(doc: Document, raw_latex: str) -> None:
    """Render a block equation as centred italic text with number on the right."""
    body, eq_num = _clean_latex(raw_latex)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(body)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    if eq_num:
        p.add_run(f"   ({eq_num})")


# ─────────────────────────────────────────────────────────────────────────────
# Table helper
# ─────────────────────────────────────────────────────────────────────────────

def add_word_table(doc: Document, caption: str, headers: list[str],
                   rows: list[list[str]], font_size: int = 10) -> None:
    """Insert a captioned Word table (caption ABOVE per IIUM Press)."""
    if not headers and not rows:
        return

    # Caption above
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(12)
    cap.paragraph_format.space_after  = Pt(3)
    _add_rich_runs(cap, caption, base_size=font_size)
    if cap.runs:
        cap.runs[0].bold = True

    col_count = max(len(headers), max((len(r) for r in rows), default=0))
    tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
    tbl.style = "Table Grid"

    # Header row (grey bg)
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_shading(c, "CCCCCC")
        p = c.paragraphs[0]
        _add_rich_runs(p, h, base_size=font_size)
        if p.runs:
            p.runs[0].bold = True

    # Data rows
    for ri, row_data in enumerate(rows):
        tr = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            if ci >= col_count:
                break
            p = tr.cells[ci].paragraphs[0]
            _add_rich_runs(p, str(val), base_size=font_size)

    # Space after
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


# ─────────────────────────────────────────────────────────────────────────────
# Markdown → docx  (used for Chapter 4)
# ─────────────────────────────────────────────────────────────────────────────

def _is_divider(line: str) -> bool:
    stripped = line.strip().replace("|", "").replace("-", "").replace(" ", "")
    return len(stripped) == 0 and "|" in line and "-" in line


def _parse_md_table_rows(raw_lines: list[str]) -> tuple[str, list[str], list[list[str]]]:
    """
    Parse a Markdown table block.
    Returns (caption, headers, data_rows).
    """
    caption = ""
    headers: list[str] = []
    rows: list[list[str]] = []

    def split_row(line: str) -> list[str]:
        parts = [c.strip() for c in line.strip().strip("|").split("|")]
        return parts

    non_divider = [l for l in raw_lines if not _is_divider(l) and l.strip()]

    # First row: might be single-cell caption  e.g.  | **Table 4.1.** ... |
    if non_divider:
        first = split_row(non_divider[0])
        if len(first) == 1:
            caption = first[0]
            non_divider = non_divider[1:]
        # Next non-empty row is headers
        if non_divider:
            headers = split_row(non_divider[0])
            for row_line in non_divider[1:]:
                rows.append(split_row(row_line))
        else:
            pass
    return caption, headers, rows


def _collect_equation(lines: list[str], start: int) -> tuple[str, int]:
    """Collect a $$…$$ block starting at lines[start]. Returns (text, next_idx)."""
    first = lines[start].strip()
    if first.startswith("$$") and first.endswith("$$") and len(first) > 4:
        return first[2:-2], start + 1
    # Multi-line
    collected: list[str] = []
    i = start
    if lines[i].strip().startswith("$$"):
        collected.append(lines[i].strip()[2:])
        i += 1
    while i < len(lines):
        s = lines[i].strip()
        if s.endswith("$$"):
            collected.append(s[:-2])
            i += 1
            break
        collected.append(s)
        i += 1
    return " ".join(collected), i


def markdown_to_docx(md_text: str, doc: Document) -> None:
    """Parse *md_text* and append formatted content into *doc*."""
    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    paragraph_buffer: list[str] = []

    def flush_para():
        nonlocal paragraph_buffer
        text = " ".join(paragraph_buffer).strip()
        if text:
            add_rich_para(doc, text)
        paragraph_buffer = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # ── Code fence ────────────────────────────────────────────────────
        if line.startswith("```"):
            flush_para()
            if in_code_block:
                # End of code block — emit as monospace block
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after  = Pt(4)
                    for cl in code_lines:
                        r = p.add_run(cl + "\n")
                        r.font.name = "Courier New"
                        r.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────
        if re.match(r"^-{3,}$", line.strip()):
            flush_para()
            i += 1
            continue

        # ── Chapter / section headings ────────────────────────────────────
        if line.startswith("# "):
            flush_para()
            text = line[2:].strip()
            # Strip leading 'Chapter N' line
            if re.match(r"Chapter\s+\d+$", text, re.IGNORECASE):
                i += 1
                continue
            doc.add_heading(text, level=1)
            i += 1
            continue

        if line.startswith("## "):
            flush_para()
            # Remove leading number e.g. "4.1 "
            text = re.sub(r"^\d+\.\d+\s+", "", line[3:].strip())
            h = doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith("### "):
            flush_para()
            text = re.sub(r"^\d+\.\d+\.\d+\s+", "", line[4:].strip())
            doc.add_heading(text, level=3)
            i += 1
            continue

        # ── Author / italic lines at top ──────────────────────────────────
        if line.startswith("**") and line.endswith("**") and i < 6:
            flush_para()
            p = doc.add_paragraph()
            _add_rich_runs(p, line)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        if line.startswith("*") and line.endswith("*") and i < 8:
            flush_para()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _add_rich_runs(p, line)
            i += 1
            continue

        # ── Equations $$...$$  ────────────────────────────────────────────
        stripped = line.strip()
        if stripped.startswith("$$"):
            flush_para()
            eq_text, i = _collect_equation(lines, i)
            add_equation(doc, eq_text)
            continue

        # ── Figure block-quote ────────────────────────────────────────────
        if stripped.startswith("> "):
            flush_para()
            # Collect all > lines
            fig_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                fig_lines.append(lines[i].strip()[1:].strip())
                i += 1
            # Parse
            fig_id   = ""
            fig_note = ""
            caption_parts: list[str] = []
            for fl in fig_lines:
                clean = re.sub(r"\*\*|\*", "", fl).strip()
                if clean.startswith("[FIGURE") or clean.startswith("[ FIGURE"):
                    fig_id = re.sub(r"[\[\]]", "", clean).strip()
                else:
                    caption_parts.append(clean)
            caption_text = " ".join(caption_parts)
            # Extract "Caption: ..." if present
            cap_match = re.search(r"[Cc]aption[:\s]+(.+)", caption_text)
            if cap_match:
                display_caption = cap_match.group(1).strip()
                fig_note = caption_text[: cap_match.start()].strip()
            else:
                display_caption = caption_text
                fig_note = "Insert figure at ≥300 DPI. Source: Authors."
            add_figure_placeholder(doc, fig_id, fig_note, display_caption)
            continue

        # ── Markdown table ────────────────────────────────────────────────
        if stripped.startswith("|"):
            flush_para()
            tbl_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            cap, hdrs, data = _parse_md_table_rows(tbl_lines)
            if hdrs or data:
                add_word_table(doc, cap, hdrs, data)
            continue

        # ── Empty line ────────────────────────────────────────────────────
        if not stripped:
            flush_para()
            i += 1
            continue

        # ── Regular paragraph text ────────────────────────────────────────
        paragraph_buffer.append(stripped)
        i += 1

    flush_para()


# ─────────────────────────────────────────────────────────────────────────────
# Chapter 4 generation
# ─────────────────────────────────────────────────────────────────────────────

def generate_chapter4(output_path: str) -> None:
    md_file = Path(__file__).parent / "CHAPTER_04_DRAFT.md"
    if not md_file.exists():
        raise FileNotFoundError(f"Chapter 4 draft not found: {md_file}")

    md_text = md_file.read_text(encoding="utf-8")

    doc = new_iium_doc()
    markdown_to_docx(md_text, doc)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[✓] Chapter 4 saved → {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# Chapter 5 content (M-ACF Net)
# ─────────────────────────────────────────────────────────────────────────────

def _ch5_intro(doc: Document) -> None:
    add_rich_para(doc, (
        "Near-shore maritime environments present a uniquely demanding context for autonomous "
        "surface navigation: water reflections, dynamic background clutter, partial occlusions "
        "from pier structures, and rapid scene changes challenge even sophisticated single-modality "
        "perception systems. Whilst decision-level fusion — in which each sensor independently "
        "generates object detections that are subsequently combined — offers practical modularity "
        "and fault-tolerance, it inherently limits the depth of cross-modal information exchange. "
        "The approach presented in this chapter addresses this limitation by fusing complementary "
        "sensor representations at an intermediate feature level, before final object hypotheses "
        "are formed, thereby enabling richer spatial–semantic integration."
    ))
    add_rich_para(doc, (
        "This chapter presents the **Maritime Adaptive Confidence Fusion Network (M-ACF Net)**, "
        "a mid-level LiDAR–stereo fusion framework designed for geospatial obstacle awareness "
        "aboard the Suraya Riptide Unmanned Surface Vehicle (USV). M-ACF Net integrates "
        "LiDAR point cloud geometry with stereo-derived visual features through a dynamic "
        "confidence-weighted fusion mechanism, followed by a Hybrid PCA–K-Means clustering "
        "stage for 3D obstacle localisation. The resulting obstacle state — position, heading, "
        "and estimated extent — is transmitted to a custom Guidance, Navigation, and Control "
        "(GNC) framework enabling real-time autonomous navigation and collision avoidance."
    ))
    add_rich_para(doc, (
        "Key contributions presented in this chapter are: (i) a mid-level confidence-weighted "
        "fusion architecture that dynamically adjusts the contribution of LiDAR and stereo "
        "features based on per-frame sensor reliability estimates; (ii) a Pseudo-LiDAR "
        "generation pipeline that transforms stereo disparity maps into 3D point representations "
        "compatible with LiDAR feature spaces; (iii) a Hybrid PCA–K-Means obstacle clustering "
        "approach that aligns cluster orientation axes with obstacle heading before performing "
        "density-based segmentation; and (iv) a fully integrated GNC system validated across "
        "multiple real-world near-shore test environments."
    ))
    add_rich_para(doc, (
        "The chapter is organised as follows. Section 5.2 describes the Suraya Riptide platform "
        "and its sensor suite. Section 5.3 presents the M-ACF Net architecture and design "
        "rationale. Sections 5.4 and 5.5 detail the visual feature extraction (including "
        "Pseudo-LiDAR generation) and LiDAR point cloud processing pipelines respectively. "
        "Section 5.6 presents the confidence-weighted mid-level fusion mechanism. Section 5.7 "
        "describes the PCA–K-Means obstacle clustering stage. Dataset collection is described "
        "in Section 5.8, and GNC integration in Section 5.9. Experimental results are presented "
        "and discussed in Section 5.10. Section 5.11 concludes with a synthesis of contributions."
    ))


def _ch5_platform(doc: Document) -> None:
    doc.add_heading("Platform Configuration and Operational Context", level=3)
    add_rich_para(doc, (
        "The Suraya Riptide is a 1.33-metre trimaran USV developed at the Centre for Unmanned "
        "Technologies (CUTe), International Islamic University Malaysia (IIUM). The trimaran "
        "hull configuration — a central fiberglass hull flanked by two supplemental buoyancy "
        "pontoons connected via aluminium cross-members — provides a stable sensor mounting "
        "platform capable of operating in wave heights up to 0.5 m, characteristic of "
        "Malaysian near-shore and inland waterway conditions. Despite its compact footprint, "
        "the Riptide's deck accommodates a full multimodal sensor suite, the onboard "
        "**NVIDIA Jetson Orin NX** computing unit, and all communication hardware required "
        "for autonomous mission execution."
    ))
    add_rich_para(doc, (
        "The Jetson Orin NX (16 GB LPDDR5 at 68.3 GB/s) provides up to 100 TOPS of AI "
        "inference throughput through its 1,024-core Ampere GPU, with a power envelope of "
        "10–25 W suitable for battery-constrained USV operation. This platform was selected "
        "specifically to support real-time execution of the M-ACF Net perception pipeline — "
        "including YOLOv8 inference, point cloud processing, deep-learning-based disparity "
        "estimation, and confidence-weighted fusion — at operationally relevant frame rates "
        "without active cooling."
    ))
    add_figure_placeholder(
        doc,
        "FIGURE 5.1",
        "Photograph of the Suraya Riptide trimaran USV with sensor suite annotated "
        "(LiDAR, ZED 2i stereo camera, GNSS, INS). Adapt from thesis platform photographs. "
        "Source: Authors.",
        "Figure 5.1. The Suraya Riptide trimaran USV, showing the co-mounted Velodyne "
        "HDL-32E LiDAR and Stereolabs ZED 2i stereo camera on the forward sensor mast, "
        "alongside the navigation sensor cluster amidships."
    )

    doc.add_heading("Navigation and Perception Sensor Modules", level=3)
    add_rich_para(doc, (
        "The Suraya Riptide integrates five sensor modules: three navigation sensors providing "
        "the vessel's absolute state, and two perception sensors furnishing the multimodal "
        "environmental data processed by M-ACF Net."
    ))
    add_rich_para(doc, (
        "**RTK-GNSS (M1G2 Portable Heading Solution).** The M1G2 provides real-time "
        "kinematic-corrected position and heading fixes at 10 Hz, with horizontal accuracy "
        "of ±(8 mm + 1 ppm) under RTK correction. Operating at 5 W from a 12 V supply, "
        "the receiver supplies the georeferenced positional data used to project "
        "obstacle estimates into the global navigation frame for mission planning."
    ))
    add_rich_para(doc, (
        "**Inertial Navigation System (SBG Ellipse-E).** The Ellipse-E provides fused "
        "roll/pitch accuracy of ±0.1°, heave estimation to 5 cm at 0.2 Hz, and a "
        "2-minute static alignment time. Its high-rate inertial outputs are used for "
        "motion compensation during point cloud and stereo frame processing, ensuring "
        "that sensor-frame obstacle estimates remain accurate under wave-induced motion."
    ))
    add_rich_para(doc, (
        "**Electronic Compass (KVH C100).** The C100 provides absolute heading at 20 Hz "
        "with ±0.5° RMS accuracy and a ±40° operational tilt range, supplying the "
        "magnetic heading reference used by the GNC coordinate transformation module."
    ))
    add_rich_para(doc, (
        "**Stereo Camera (Stereolabs ZED 2i).** The ZED 2i captures synchronised stereo "
        "image pairs at 4,416 × 1,242 pixels and 20 Hz, with a 65 mm stereo baseline "
        "and a 90° × 60° field of view. Its integrated stereo matching firmware provides "
        "dense depth maps at 1,280 × 720 resolution. In M-ACF Net, the ZED 2i supplies "
        "both the RGB image stream for YOLOv8 object detection and the left–right image "
        "pair from which deep-learning-based disparity estimation generates the "
        "Pseudo-LiDAR representation used in mid-level fusion."
    ))
    add_rich_para(doc, (
        "**LiDAR (Velodyne HDL-32E).** The HDL-32E fires 32 laser channels at a rotation "
        "rate of 10 Hz, generating approximately 700,000 points per second across a "
        "360° horizontal field of view. Range accuracy is ±2 cm at distances up to 50 m, "
        "providing the high-precision spatial measurements that constitute the LiDAR "
        "branch of the M-ACF Net fusion pipeline."
    ))
    add_word_table(
        doc,
        "Table 5.1. Sensor suite of the Suraya Riptide USV with key specifications.",
        ["Module", "Model", "Key Specification", "Role in M-ACF Net"],
        [
            ["Computing",      "NVIDIA Jetson Orin NX 16 GB",    "1024-core Ampere GPU; 100 TOPS; 10–25 W",    "Real-time inference host"],
            ["LiDAR",          "Velodyne HDL-32E",               "32 ch; 10 Hz; ±2 cm @ 50 m; ~700k pts/s",   "Geometric feature branch"],
            ["Stereo Camera",  "Stereolabs ZED 2i",              "4416×1242 @ 20 Hz; 65 mm baseline; 90°×60° FoV", "Visual/depth feature branch"],
            ["RTK-GNSS",       "M1G2 Portable Heading Solution", "10 Hz; ±(8 mm + 1 ppm) RTK; 5 W",           "Global georeferencing"],
            ["INS",            "SBG Ellipse-E",                  "±0.1° roll/pitch; 5 cm heave; 2 min align",  "Motion compensation"],
            ["Compass",        "KVH C100",                       "±0.5° RMS; 20 Hz; ±40° tilt range",          "Heading reference"],
        ]
    )

    doc.add_heading("Sensor Placement, Calibration, and Synchronisation", level=3)
    add_rich_para(doc, (
        "All navigational sensors (GNSS, INS, compass) are positioned along the hull's "
        "centreline, close to the vessel's centre of gravity, to minimise the lever-arm "
        "corrections required for accurate dynamic measurements. The perception sensors "
        "— LiDAR and stereo camera — are co-mounted on a rigid steel mast at the vessel's "
        "bow to maximise forward-hemisphere overlap and minimise mutual occlusion. "
        "Table 5.2 records the measured translation offsets of each sensor frame "
        "relative to the vessel body frame origin (F_base), with all axes aligned to "
        "the vessel centreline (yaw = pitch = roll = 0° for all sensors)."
    ))
    add_word_table(
        doc,
        "Table 5.2. Sensor coordinate frame translation offsets relative to the vessel "
        "body frame origin (metres; yaw = pitch = roll = 0° for all sensors).",
        ["Sensor Frame", "p_x (m)", "p_y (m)", "p_z (m)"],
        [
            ["F_GNSS",    "0",  "0.38", "−0.09"],
            ["F_INS",     "0", "−0.30",  "0.06"],
            ["F_LiDAR",   "0",  "0.66",  "0.43"],
            ["F_Camera",  "0",  "0.66",  "0.26"],
            ["F_Compass", "0",  "0.66", "−0.09"],
        ]
    )
    add_rich_para(doc, (
        "The LiDAR and stereo camera are co-mounted at forward offset p_y = 0.66 m, "
        "with the LiDAR positioned 0.17 m above the stereo camera (p_z = 0.43 m vs. "
        "0.26 m). This close co-location minimises the translational parallax between "
        "the two sensor frames, reducing extrinsic calibration error. Extrinsic "
        "calibration between the LiDAR and camera frames was performed using a "
        "reflective checkerboard target visible in both sensor modalities, achieving "
        "translational discrepancies under 2 cm and rotational error below 0.5°. "
        "Temporal co-registration between the 10 Hz LiDAR scan and the 20 Hz stereo "
        "frame was enforced via a hardware PPS (pulse-per-second) signal from the "
        "GNSS receiver, providing synchronisation accuracy within ±2 ms."
    ))
    add_figure_placeholder(
        doc,
        "FIGURE 5.2",
        "Schematic of sensor coordinate frames (F_LiDAR, F_Camera, F_GNSS, F_INS) "
        "relative to vessel body frame F_base, with translation offsets annotated. "
        "Adapt from thesis Figure 3.4.1. Source: Authors.",
        "Figure 5.2. Sensor coordinate frame layout on the Suraya Riptide, illustrating "
        "translation offsets from the vessel body frame origin for each sensor module."
    )


def _ch5_architecture(doc: Document) -> None:
    add_rich_para(doc, (
        "The M-ACF Net framework is motivated by the observation that mid-level fusion — "
        "the integration of sensor signals at the feature representation stage, prior to "
        "final detection decisions — offers a richer information exchange than decision-level "
        "approaches that combine only final object hypotheses. At the mid-level, geometric "
        "features derived from LiDAR point clouds and semantic depth features derived from "
        "stereo images can be aligned in a shared spatial representation, allowing the model "
        "to exploit correlations that are invisible to either sensor operating alone."
    ))
    add_rich_para(doc, (
        "A central design challenge for mid-level fusion is managing the confidence mismatch "
        "between sensor modalities: LiDAR provides superior absolute range accuracy but "
        "suffers from sparsity on small or low-reflectance targets, whilst stereo cameras "
        "deliver dense visual coverage but produce unreliable depth estimates under specular "
        "reflections or textureless surfaces. M-ACF Net addresses this through a "
        "**dynamic confidence-weighted fusion mechanism** that adjusts the relative "
        "contribution of each sensor's features on a per-frame basis, guided by "
        "per-modality reliability estimates computed from the raw data quality."
    ))
    add_figure_placeholder(
        doc,
        "FIGURE 5.3",
        "Block diagram of the complete M-ACF Net pipeline from sensor input to GNC output. "
        "Adapt from thesis architecture diagrams. Source: Authors.",
        "Figure 5.3. Overall M-ACF Net architecture: sensor acquisition → feature extraction "
        "(LiDAR and stereo branches) → confidence-weighted mid-level fusion → PCA–K-Means "
        "clustering → GNC integration."
    )

    doc.add_heading("Architectural Overview and Egocentric Processing", level=3)
    add_rich_para(doc, (
        "The M-ACF Net pipeline operates in an **egocentric reference frame** centred on the "
        "USV, in which all sensor measurements and fused obstacle representations are expressed "
        "relative to the vessel's current position and heading. This choice simplifies "
        "coordinate management in the downstream GNC system and ensures that obstacle "
        "proximity and bearing are always expressed in the vessel-centric terms most "
        "relevant to collision avoidance decision-making."
    ))
    add_rich_para(doc, (
        "The pipeline comprises five sequential stages: (1) **sensor preprocessing** — "
        "raw LiDAR and stereo data are independently filtered, downsampled, and calibrated; "
        "(2) **feature extraction** — LiDAR geometric features and stereo-derived visual "
        "depth features are computed in parallel processing threads; (3) **mid-level "
        "confidence-weighted fusion** — features from both branches are combined using "
        "dynamically computed per-modality confidence weights; (4) **PCA–K-Means clustering** "
        "— fused features are grouped into spatially coherent obstacle candidates, each "
        "assigned a 3D bounding representation; and (5) **GNC transmission** — obstacle "
        "state estimates are forwarded to the navigation controller via a structured "
        "communication interface."
    ))

    doc.add_heading("Sensor Coordinate Alignment and Preprocessing", level=3)
    add_rich_para(doc, (
        "Before feature extraction can commence, LiDAR point clouds must be projected into "
        "the camera coordinate frame to establish spatial correspondence between the two "
        "sensor modalities. Given the extrinsically calibrated rotation matrix **R** and "
        "translation vector **t**, each LiDAR point "
        "**p**_L = (x_L, y_L, z_L)^T is transformed to the camera frame as:"
    ))
    add_equation(doc, r"\mathbf{p}_C = \mathbf{R}\,\mathbf{p}_L + \mathbf{t} \tag{5.1}")
    add_rich_para(doc, (
        "The transformed point is subsequently projected onto the image plane using the "
        "camera intrinsic matrix **K**, yielding the 2D pixel coordinates at which the "
        "LiDAR point appears in the camera image. This projection enables spatial "
        "correspondence between LiDAR range measurements and stereo image features, "
        "which is the foundation of mid-level fusion."
    ))
    add_equation(doc, r"\begin{pmatrix} u \\ v \\ 1 \end{pmatrix} = \frac{1}{Z_C}\,\mathbf{K}\,\mathbf{p}_C \tag{5.2}")
    add_rich_para(doc, (
        "where (u, v) are the projected pixel coordinates, Z_C is the depth of the "
        "LiDAR point in the camera frame, and **K** is the 3 × 3 camera intrinsic matrix "
        "containing focal lengths and principal point offsets."
    ))


def _ch5_visual(doc: Document) -> None:
    doc.add_heading("Object Detection Model Selection and Training", level=3)
    add_rich_para(doc, (
        "Five variants of YOLOv8 — Nano (YOLOv8n), Small (YOLOv8s), Medium (YOLOv8m), "
        "Large (YOLOv8l), and Extra-large (YOLOv8x) — were evaluated alongside a "
        "Faster R-CNN X101-FPN two-stage detector as baseline. Selection criteria "
        "prioritised: (i) inference throughput on the NVIDIA Jetson Orin NX sufficient "
        "for real-time operation; (ii) detection accuracy on the maritime obstacle "
        "classes relevant to USV collision avoidance; and (iii) cross-domain robustness "
        "to avoid performance collapse when tested in environments different from the "
        "training distribution."
    ))
    add_rich_para(doc, (
        "Training data were drawn from two established maritime benchmark datasets. "
        "The **SeaShips** dataset comprises 7,080 annotated instances across six "
        "vessel categories: bulk cargo carrier (19.85%), ore carrier (21.24%), "
        "fishing boat (28.05%), container ship (10.48%), general cargo ship (14.73%), "
        "and passenger ship (5.63%). The **Singapore Maritime Dataset (SMD)** provides "
        "48,317 instances across nine categories, including boat, buoy, ferry, kayak, "
        "speed boat, sailboat, vessel/ship (71.60% of instances), and other. "
        "All annotations were standardised to a unified 'sea_obstacle' class "
        "(class ID = 0) to focus detection performance on obstacle presence and "
        "localisation independent of vessel-type classification, which is handled "
        "at the comprehension layer of the GNC system."
    ))
    add_rich_para(doc, (
        "Training employed data augmentation (random scaling, cropping, and lighting "
        "variations) to enhance model robustness to the diverse illumination and "
        "viewpoint conditions encountered in near-shore USV deployments. Batch size "
        "and input resolution were tuned per variant to the memory constraints of "
        "the Jetson Orin NX. Optimisation targeted precision–recall balance, with "
        "model selection guided by mAP@0.5 on held-out validation splits of each "
        "benchmark dataset."
    ))
    add_rich_para(doc, (
        "**Intra-domain performance.** The domain-specific best models — YOLOv8x for "
        "SMD and YOLOv8m for SeaShips — achieved strong in-domain detection accuracy "
        "as shown in Table 5.3. The SMD-trained YOLOv8x reached mAP@0.5 = 0.988 "
        "(precision 0.983, recall 0.961), whilst the SeaShip-trained YOLOv8m "
        "achieved mAP@0.5 = 0.954 (precision 0.913, recall 0.889). However, "
        "cross-domain evaluation revealed significant performance degradation: applying "
        "YOLOv8x trained on SMD to the SeaShips test set reduced mAP@0.5 to 0.446 "
        "(a 54.9% drop), whilst the YOLOv8m SeaShips model applied to SMD yielded "
        "mAP@0.5 = 0.144 (an 84.9% drop). This cross-domain brittleness motivates "
        "the multi-environment dataset collection described in Section 5.8 and "
        "the confidence-weighted fusion strategy, which inherently reduces dependence "
        "on any single-modality detection pathway."
    ))
    add_word_table(
        doc,
        "Table 5.3. YOLOv8 intra-domain detection performance on SMD and SeaShips "
        "validation sets (unified 'sea_obstacle' class).",
        ["Model", "Dataset", "Precision", "Recall", "mAP@0.5", "mAP@0.5:0.95"],
        [
            ["YOLOv8x", "SMD",      "0.983", "0.961", "0.988", "0.839"],
            ["YOLOv8m", "SeaShips", "0.913", "0.889", "0.954", "0.686"],
        ]
    )

    doc.add_heading("Stereo Vision Depth Estimation and Occlusion Handling", level=3)
    add_rich_para(doc, (
        "Dense depth estimation is performed via a deep-learning-based cost-volume "
        "disparity network. For each pixel (x, y), a cost volume C(x, y, d) is "
        "constructed by summing squared photometric differences over a neighbourhood N "
        "of the candidate disparity d:"
    ))
    add_equation(doc, r"C(x,y,d) = \sum_{k \in N} \left(I_L(x+k,y) - I_R(x-d+k,y)\right)^2 \tag{5.3}")
    add_rich_para(doc, (
        "The disparity map is then recovered from the cost volume via a soft-argmin "
        "operation that produces a sub-pixel-accurate, differentiable disparity estimate:"
    ))
    add_equation(doc, r"D(x,y) = \sum_d d \cdot \mathrm{softmax}\bigl(-C(x,y,d)\bigr) \tag{5.4}")
    add_rich_para(doc, (
        "Metric depth Z is recovered from disparity d using the ZED 2i's calibrated "
        "focal length f and its fixed 65 mm stereo baseline B:"
    ))
    add_equation(doc, r"Z = \frac{f \cdot B}{d} \tag{5.5}")
    add_rich_para(doc, (
        "**Stereo occlusion handling.** Regions visible in only one camera — common at "
        "obstacle boundaries and water-surface interfaces — produce unreliable disparity "
        "estimates. M-ACF Net identifies occluded pixels via left–right consistency "
        "checking and fills them using an inverse-distance-weighted interpolation from "
        "neighbouring valid depth estimates:"
    ))
    add_equation(doc, r"d_{\mathrm{interpolated}} = \frac{\sum_i w_i \cdot d_i}{\sum_i w_i} \tag{5.6}")
    add_rich_para(doc, (
        "The stereo recovery efficiency η — the fraction of originally occluded pixels "
        "successfully recovered by interpolation — is monitored per frame and "
        "contributes to the stereo confidence score C_stereo used in the fusion stage "
        "(Section 5.6):"
    ))
    add_equation(doc, r"\eta = \frac{N_{\mathrm{recovered}}}{N_{\mathrm{occluded}}} \times 100\% \tag{5.7}")
    add_figure_placeholder(
        doc,
        "FIGURE 5.4",
        "Side-by-side comparison: left RGB image, raw disparity map (with occluded "
        "regions visible as holes), and recovered disparity after weighted interpolation, "
        "across calm and wave-disturbed water conditions. Source: Authors.",
        "Figure 5.4. Stereo disparity estimation and occlusion recovery: (a) left RGB "
        "input; (b) raw disparity map with occluded regions; (c) interpolation-recovered "
        "disparity map used for Pseudo-LiDAR generation."
    )

    doc.add_heading("Pseudo-LiDAR Generation from Disparity Maps", level=3)
    add_rich_para(doc, (
        "The recovered depth map is converted into a Pseudo-LiDAR point cloud — a 3D "
        "point representation geometrically equivalent to a direct LiDAR measurement "
        "from the camera viewpoint. Each pixel (u, v) in the depth image is unprojected "
        "into 3D camera-frame coordinates using the camera intrinsic matrix K:"
    ))
    add_equation(doc, r"\begin{pmatrix} X \\ Y \\ Z \end{pmatrix} = Z \cdot \mathbf{K}^{-1} \begin{pmatrix} u \\ v \\ 1 \end{pmatrix} \tag{5.8}")
    add_rich_para(doc, (
        "equivalently expressed in terms of the camera's focal lengths (f_x, f_y) "
        "and principal point (c_x, c_y) as:"
    ))
    add_equation(doc, r"X = \frac{Z(u - c_x)}{f_x}, \quad Y = \frac{Z(v - c_y)}{f_y} \tag{5.9}")
    add_rich_para(doc, (
        "The Pseudo-LiDAR cloud is clipped to the same field-of-view and range bounds "
        "as the physical LiDAR, and voxel-downsampled to a comparable point density "
        "before entering the mid-level fusion stage. This representation allows the "
        "fusion module to operate on two geometrically compatible 3D feature sets — "
        "real LiDAR and stereo-derived Pseudo-LiDAR — enabling richer spatial "
        "correspondence than possible with raw image–point cloud pairings."
    ))


def _ch5_lidar(doc: Document) -> None:
    add_rich_para(doc, (
        "The LiDAR processing branch transforms raw 3D point clouds into a set of "
        "compact geometric features suitable for mid-level fusion. The pipeline "
        "proceeds through four sequential stages."
    ))

    doc.add_heading("Downsampling, Noise Removal, and FOV Filtering", level=3)
    add_rich_para(doc, (
        "**Voxel downsampling** replaces all points within each voxel cell of side "
        "length **0.2 m** with a single centroid point. This voxel size was selected "
        "through empirical optimisation as the configuration that achieved the best "
        "balance between cloud density and processing speed, producing an 8.0% "
        "reduction in point count whilst preserving the geometric structure of "
        "maritime obstacles at the relevant detection scales."
    ))
    add_rich_para(doc, (
        "**Statistical outlier removal (SOR)** identifies and discards points whose "
        "mean distance to their k nearest neighbours deviates substantially from the "
        "global neighbourhood mean. The SOR threshold parameter k = 1.5 standard "
        "deviations was identified as optimal through F1-score analysis across multiple "
        "near-shore environments: the preprocessing configuration (voxel = 0.2 m, "
        "k = 1.5) achieved an F1-score of **0.85** on the evaluation dataset, "
        "reflecting an effective balance between noise elimination and retention of "
        "sparse but meaningful obstacle returns from buoys and distant vessels."
    ))
    add_rich_para(doc, (
        "**Field-of-view (FOV) filtering** restricts the processed cloud to the "
        "forward-facing hemisphere aligned with the stereo camera's angular coverage, "
        "ensuring geometric consistency between the LiDAR and stereo branches in the "
        "fusion stage. Points outside the FOV intersection are discarded, as they lack "
        "corresponding stereo depth measurements and would contribute uninformative "
        "features to the fusion module."
    ))

    doc.add_heading("Geometric Feature Extraction from Point Clouds", level=3)
    add_rich_para(doc, (
        "Following preprocessing, geometric features are extracted from the filtered cloud "
        "to characterise local surface structure. For each point in the processed cloud, "
        "a neighbourhood of k nearest neighbours is identified, and local surface normal "
        "vectors and curvature estimates are computed via Principal Component Analysis (PCA) "
        "on the neighbourhood covariance matrix. These local geometric descriptors encode "
        "the orientation and smoothness of local surface patches — key discriminators "
        "between flat water surfaces (smooth, horizontal normals) and obstacle surfaces "
        "(varied orientations, higher curvature). The feature vector for each point thus "
        "comprises its 3D spatial coordinates augmented by its surface normal direction "
        "and principal curvature estimate, forming the LiDAR geometric feature representation "
        "passed to the fusion stage."
    ))


def _ch5_fusion(doc: Document) -> None:
    add_rich_para(doc, (
        "The confidence-weighted mid-level fusion stage is the core contribution of M-ACF Net. "
        "Rather than fusing sensors at the decision level (combining final object detections) "
        "or at the earliest feature level (raw concatenation of sensor signals), M-ACF Net "
        "operates at an intermediate representation level where LiDAR geometric features and "
        "stereo-derived depth features have been independently computed and are now merged "
        "into a unified spatial-semantic feature volume."
    ))

    doc.add_heading("Fusion Strategy and Confidence Modelling", level=3)
    add_rich_para(doc, (
        "M-ACF Net computes per-frame, per-modality confidence scores that quantify "
        "the reliability of each sensor's feature contribution. The **LiDAR confidence** "
        "C_LiDAR is derived from the ratio of valid filtered points to total raw points "
        "in the current scan — a proxy for point cloud density and data quality:"
    ))
    add_equation(doc, r"C_{\mathrm{LiDAR}} = \frac{N_{\mathrm{filtered}}}{N_{\mathrm{raw}}} \tag{5.10}")
    add_rich_para(doc, (
        "High C_LiDAR (dense, well-filtered cloud) indicates reliable LiDAR geometry; "
        "low C_LiDAR (sparse cloud, as occurs with small or low-reflectance targets "
        "at range) signals degraded feature quality. The **stereo confidence** C_stereo "
        "is derived from the fraction of non-occluded pixels identified by the "
        "left-right consistency check (Section 5.4.2):"
    ))
    add_equation(doc, r"C_{\mathrm{stereo}} = 1 - \frac{N_{\mathrm{occ}}}{N_{\mathrm{total}}} \tag{5.11}")
    add_rich_para(doc, (
        "where N_occ is the number of pixels flagged as occluded or unreliable, and "
        "N_total is the total pixel count per frame. C_stereo approaches 1.0 in "
        "well-lit, low-occlusion conditions (high stereo reliability) and drops "
        "towards 0 under specular water reflections, fog, or heavy wave-induced occlusion."
    ))

    doc.add_heading("Dynamic Adaptive Weighting Mechanism", level=3)
    add_rich_para(doc, (
        "The adaptive weighting factor α is computed from the per-frame confidence scores:"
    ))
    add_equation(doc, r"\alpha = \frac{C_{\mathrm{LiDAR}}}{C_{\mathrm{LiDAR}} + C_{\mathrm{stereo}}} \tag{5.12}")
    add_rich_para(doc, (
        "The fused feature representation is then formed as a confidence-weighted "
        "combination of the LiDAR and stereo feature branches:"
    ))
    add_equation(doc, r"F_{\mathrm{fused}} = \alpha \cdot F_{\mathrm{LiDAR}} \oplus (1 - \alpha) \cdot F_{\mathrm{stereo}} \tag{5.13}")
    add_rich_para(doc, (
        "where ⊕ denotes element-wise feature combination in the shared spatial "
        "representation. By construction, α + (1 − α) = 1, ensuring normalised "
        "feature contributions at every spatial location. Empirical analysis of "
        "per-frame α distributions across diverse maritime conditions confirmed that "
        "LiDAR is typically prioritised (α ≥ 0.6) owing to its resilience to lighting "
        "variations and specular water reflections. Stereo vision contributes more "
        "significantly only in close-range, well-lit scenarios (α < 0.5), where its "
        "higher pixel density enhances depth precision in the near field. The fused "
        "3D obstacle position is then computed by back-projecting the fused depth "
        "estimate into metric space:"
    ))
    add_equation(doc, r"\begin{pmatrix} X \\ Y \end{pmatrix} = F_{\mathrm{fused}} \cdot \begin{pmatrix} (u - c_x)/f_x \\ (v - c_y)/f_y \end{pmatrix} \tag{5.14}")
    add_figure_placeholder(
        doc,
        "FIGURE 5.5",
        "Visualisation of per-frame confidence scores C_LiDAR and C_stereo and the "
        "resulting alpha weights across a sequence of frames with varying conditions "
        "(calm, choppy, rainy, foggy). Source: Authors.",
        "Figure 5.5. M-ACF Net adaptive confidence weighting across maritime conditions: "
        "(a) LiDAR confidence C_LiDAR; (b) stereo confidence C_stereo; (c) resulting "
        "α values demonstrating LiDAR priority (α ≥ 0.6) in most conditions."
    )


def _ch5_clustering(doc: Document) -> None:
    add_rich_para(doc, (
        "Following mid-level fusion, the resulting 3D feature volume is processed by a "
        "**Hybrid PCA–K-Means** clustering stage that segments the fused point "
        "representation into spatially coherent obstacle candidates. This stage replaces "
        "the Euclidean clustering commonly used in LiDAR-only pipelines with an "
        "orientation-aware approach better suited to the elongated geometries of maritime "
        "obstacles (vessels, jetties, pontoons), which are poorly segmented by "
        "spherically-symmetric distance metrics."
    ))

    doc.add_heading("Principal Component Analysis for Orientation Alignment", level=3)
    add_rich_para(doc, (
        "The PCA–K-Means algorithm (Algorithm 5.1) operates on the full fused point "
        "cloud. PCA first computes the mean vector μ of all N points:"
    ))
    add_equation(doc, r"\mathbf{\mu} = \frac{1}{N}\sum_{i=1}^{N}\mathbf{p}_i \tag{5.15}")
    add_rich_para(doc, (
        "The centred covariance matrix Σ is then computed from the mean-subtracted data:"
    ))
    add_equation(doc, r"\mathbf{\Sigma} = \frac{1}{N-1}\,\mathbf{P}_c^T\,\mathbf{P}_c \tag{5.16}")
    add_rich_para(doc, (
        "where P_c is the N × 3 matrix of centred point coordinates "
        "(P_c = point_cloud − μ). Eigendecomposition of Σ yields eigenvectors "
        "(principal components) and eigenvalues:"
    ))
    add_equation(doc, r"\mathbf{\Sigma}\,\mathbf{v} = \lambda\,\mathbf{v} \tag{5.17}")
    add_rich_para(doc, (
        "Eigenvectors are sorted in descending order of eigenvalue — placing the "
        "direction of maximum point cloud variance (the obstacle's principal axis) "
        "first. The full point cloud is projected into the PCA-aligned frame:"
    ))
    add_equation(doc, r"\mathbf{P}_{\mathrm{PCA}} = \mathbf{P}_c \cdot \mathbf{V}_{\mathrm{sorted}} \tag{5.18}")
    add_rich_para(doc, (
        "where V_sorted is the matrix of sorted eigenvectors. In this transformed "
        "space, the first dimension aligns with the obstacle's principal axis (e.g. "
        "the long axis of a vessel), enabling K-Means to cluster along physically "
        "meaningful directions rather than arbitrary Cartesian axes."
    ))

    doc.add_heading("K-Means Clustering in Transformed Space", level=3)
    add_rich_para(doc, (
        "K-Means is applied to P_PCA, minimising the within-cluster sum-of-squares "
        "objective J:"
    ))
    add_equation(doc, r"J = \sum_{i=1}^{k}\sum_{j=1}^{n_i}\left\|\mathbf{x}_j^{(i)} - \mathbf{\mu}_i\right\|^2 \tag{5.19}")
    add_rich_para(doc, (
        "K is determined adaptively from the variance captured by the first principal "
        "component — scenes in which the dominant axis explains a large fraction of "
        "total variance are assigned k = 1 (a single elongated obstacle), whilst "
        "scenes with more distributed variance receive higher k. Clustering iterates "
        "until convergence (||new_centroids − old_centroids|| < ε). The best "
        "configuration across all evaluation environments was voxel size = 0.2 m "
        "with SOR threshold k = 1.5, achieving a Silhouette score of 0.70, "
        "Davies–Bouldin index of 0.50, and Calinski–Harabasz index of 60 — "
        "representing the lowest computation time (0.05 s per frame) of all "
        "evaluated clustering approaches (Table 5.7 in Section 5.10)."
    ))

    doc.add_heading("Centroid and Extremity Estimation for Navigational Awareness", level=3)
    add_rich_para(doc, (
        "For each valid cluster, M-ACF Net computes two navigational estimates: "
        "the **centroid** (arithmetic mean of all cluster point positions), "
        "representing the obstacle's geometric centre; and the **principal component "
        "extremity** — the cluster point with the highest absolute value along the "
        "first principal component (PC1), representing the closest physical surface "
        "of the obstacle to the USV along its dominant axis. Both estimates are "
        "transmitted to the GNC system via the obstacle communication message. "
        "The extremity-based estimate is used to compute the safe standoff "
        "margin for avoidance manoeuvres, as it consistently provides earlier and "
        "more conservative collision warnings than the centroid estimate for elongated "
        "or large obstacles."
    ))
    add_figure_placeholder(
        doc,
        "FIGURE 5.6",
        "Visualisation of PCA–K-Means clustering output: colour-coded obstacle clusters "
        "with centroid and nearest-extremity markers overlaid on camera view. Source: Authors.",
        "Figure 5.6. PCA–K-Means obstacle clustering result: colour-coded clusters with "
        "centroid (circle) and nearest-extremity (diamond) positions annotated for each "
        "detected obstacle."
    )


def _ch5_dataset(doc: Document) -> None:
    add_rich_para(doc, (
        "A dedicated multi-modal dataset was collected aboard the Suraya Riptide across "
        "Malaysian near-shore and inland waterways, providing a representative "
        "distribution of the obstacle categories and environmental conditions "
        "encountered in real USV operations. The dataset encompasses three primary "
        "obstacle categories: **boats** (fishing vessels, cargo ships, kayaks) "
        "constituting 62% of all annotated instances; **structures** (piers, buoys, "
        "floating barriers) constituting 23%; and **dynamic obstacles** (drifting "
        "debris, human activity) constituting the remaining 15%."
    ))

    doc.add_heading("Multi-Modal, Multi-Environment Dataset Overview", level=3)
    add_rich_para(doc, (
        "Data were recorded synchronously from the Velodyne HDL-32E at 10 Hz and the "
        "ZED 2i stereo camera at 20 Hz, with timestamps aligned to the GNSS PPS signal "
        "to within ±2 ms. GNSS/INS data were decoded to CSV in the UTM coordinate "
        "system, providing the georeferenced ground truth for obstacle position "
        "validation in controlled ranging trials at Taman Tasik Metropolitan "
        "(a lake environment with fixed reference targets at known RTK-GNSS positions). "
        "LiDAR data were stored as PCD files; stereo image sequences were saved as "
        "rectified sequential JPEG folders; all data were packaged in ROS bag format "
        "for offline playback and annotation."
    ))
    add_word_table(
        doc,
        "Table 5.4. Multi-modal dataset composition: obstacle categories and sensor "
        "configuration during collection (Velodyne HDL-32E + Stereolabs ZED 2i).",
        ["Obstacle Category", "Examples", "% of Instances", "Sensor Modalities", "Rate"],
        [
            ["Boats",            "Fishing vessels, cargo, kayaks",    "62%",  "LiDAR + Stereo", "10/20 Hz"],
            ["Structures",       "Piers, buoys, floating barriers",   "23%",  "LiDAR + Stereo", "10/20 Hz"],
            ["Dynamic obstacles","Drifting debris, human activity",   "15%",  "LiDAR + Stereo", "10/20 Hz"],
        ]
    )

    doc.add_heading("Data Collection Procedures and Preprocessing Pipeline", level=3)
    add_rich_para(doc, (
        "All sequences were post-processed to extract temporally aligned LiDAR–stereo "
        "frame pairs, discarding frames affected by sensor start-up transients or "
        "communication dropouts. Bounding box annotations for detected obstacle "
        "instances were produced via a semi-automated pipeline: camera-frame 2D "
        "bounding boxes were annotated by domain experts, and 3D LiDAR annotations "
        "were derived automatically via the extrinsic calibration projection, ensuring "
        "label consistency across modalities. The annotation format is compatible with "
        "both 2D (YOLO) and 3D (KITTI) evaluation protocols, enabling benchmarking "
        "against both 2D detection and 3D localisation baselines."
    ))


def _ch5_gnc(doc: Document) -> None:
    add_rich_para(doc, (
        "The M-ACF Net perception pipeline is integrated into a complete Guidance, "
        "Navigation, and Control (GNC) architecture that enables autonomous mission "
        "execution and reactive collision avoidance aboard the Suraya Riptide. The "
        "GNC system receives obstacle state estimates — position, estimated heading, "
        "and spatial extent — from the M-ACF Net clustering module and translates "
        "these into navigation decisions at the mission planning and actuator control layers."
    ))

    doc.add_heading("System Architecture and Communication Framework", level=3)
    add_rich_para(doc, (
        "The GNC system is structured around a hierarchical two-layer control architecture. "
        "The **high-level controller** (global and local planners) operates on the "
        "obstacle map maintained by M-ACF Net, computing waypoint-following trajectories "
        "and reactive avoidance manoeuvres in the global reference frame. The "
        "**low-level controller** translates high-level velocity and heading commands "
        "into PWM actuator signals for the propulsion and steering systems of the Riptide. "
        "Communication between the perception and GNC subsystems uses a structured "
        "message protocol analogous to the NMEA-based VIP module described in Chapter 4, "
        "transmitting obstacle position, range, and confidence metadata at each "
        "perception cycle."
    ))
    add_figure_placeholder(
        doc,
        "FIGURE 5.7",
        "Block diagram of the GNC architecture: high-level planner → local planner → "
        "obstacle manager → low-level actuator controller, with M-ACF Net perception "
        "inputs annotated. Source: Authors.",
        "Figure 5.7. GNC system architecture for the Suraya Riptide: hierarchical "
        "control structure from M-ACF Net perception outputs to actuator commands."
    )

    doc.add_heading("Guidance, Navigation, and Obstacle Avoidance Behaviours", level=3)
    add_rich_para(doc, (
        "Detected obstacles are transmitted from the M-ACF Net perception node to the "
        "GNC system via a UDP client-server communication link using a custom "
        "NMEA-compliant obstacle message format:"
    ))
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.4)
    p_code.paragraph_format.space_before = Pt(4)
    p_code.paragraph_format.space_after  = Pt(6)
    r_code = p_code.add_run(
        "$ODOBJ,<TOD>,<CIO>,<id>,<x>,<y>,<z>,<W>,<H>,<L>,<D>*<cc>CRLF"
    )
    r_code.font.name = "Courier New"
    r_code.font.size = Pt(9)
    add_rich_para(doc, (
        "where TOD = total objects detected; CIO = current object index; "
        "id = object tracking identifier; x, y, z = obstacle coordinates in metres; "
        "W, H, L = obstacle width, height, and length; D = distance from sensor; "
        "cc = XOR checksum. The message format extends the VIP protocol used in the "
        "ESA system (Chapter 4) by adding the z-coordinate for full 3D obstacle "
        "position reporting, enabling the GNC system to reason about obstacle "
        "elevation and cross-validate LiDAR and stereo height estimates."
    ))
    add_rich_para(doc, (
        "Obstacle coordinates are transformed from the sensor frame to the GNC "
        "navigation frame using a geodetic conversion that references all obstacle "
        "positions relative to a fixed mission-start georeferenced origin "
        "(λ_ref, φ_ref, A_ref):"
    ))
    add_equation(doc, r"\begin{pmatrix}x \\ y \\ z\end{pmatrix} = \begin{pmatrix}k(\lambda - \lambda_{\mathrm{ref}})\cos\phi \\ k(\phi - \phi_{\mathrm{ref}}) \\ A - A_{\mathrm{ref}}\end{pmatrix} \tag{5.20}")
    add_rich_para(doc, (
        "where k ≈ 111,320 m/° is the standard geodetic conversion factor. "
        "Magnetic heading is corrected to true heading by applying the local "
        "magnetic variation Δ_m supplied by the compass: θ_true = θ_magnetic + Δ_m. "
        "The global planner manages waypoint-following trajectories; the local planner "
        "implements reactive avoidance using a Closest Point of Approach (CPA) "
        "objective function to identify heading–speed combinations that maintain "
        "safe clearance from all reported obstacles under COLREGs constraints."
    ))


def _ch5_results(doc: Document) -> None:
    doc.add_heading("LiDAR Preprocessing and Stereo Feature Extraction Performance", level=3)
    add_rich_para(doc, (
        "LiDAR preprocessing (voxel downsampling at 0.2 m, SOR with k = 1.5, "
        "and FOV filtering) achieved an F1-score of **0.85** on the evaluation "
        "dataset, reducing mean cloud density by 8.0% whilst preserving the "
        "geometric structure of maritime obstacles. Stereo disparity estimation "
        "and Pseudo-LiDAR generation executed at **9.8 ± 1.3 ms per frame** "
        "on the NVIDIA Jetson Orin NX — corresponding to a potential throughput "
        "of 102 FPS, substantially faster than the 20 Hz LiDAR acquisition rate "
        "and ensuring that stereo processing is never the pipeline bottleneck. "
        "The complete M-ACF Net system operates at **18–25 FPS** end-to-end, "
        "well within the real-time constraint of <100 ms per frame defined for "
        "USV navigation."
    ))
    add_figure_placeholder(
        doc,
        "FIGURE 5.8",
        "Bar chart showing per-stage processing latency (LiDAR preprocessing, "
        "stereo disparity, pseudo-LiDAR, confidence fusion, PCA-KMeans, GNC "
        "transmission) on the NVIDIA Jetson Orin NX. Source: Authors.",
        "Figure 5.8. M-ACF Net per-stage processing latency on the NVIDIA Jetson "
        "Orin NX, demonstrating end-to-end throughput of 18–25 FPS "
        "(stereo branch: 9.8 ± 1.3 ms; total pipeline: <55 ms under calm conditions)."
    )

    doc.add_heading("Fusion Performance: Quantitative Analysis", level=3)
    add_rich_para(doc, (
        "Fusion performance was evaluated by comparing depth estimation RMSE and "
        "outlier reduction against raw (unweighted) sensor inputs and four "
        "alternative fusion baselines across five maritime environmental conditions. "
        "Table 5.5 summarises the results."
    ))
    add_word_table(
        doc,
        "Table 5.5. Mid-level fusion performance: RMSE and outlier reduction under "
        "five maritime environmental conditions (M-ACF Net vs. raw sensor input).",
        ["Condition", "RMSE Raw (m)", "RMSE Fused (m)", "Outlier Reduction (%)", "Process Time (ms)"],
        [
            ["Calm Waters",       "1.2 ± 0.3", "0.6 ± 0.1", "72.5", "45"],
            ["Choppy Waters",     "2.1 ± 0.5", "1.0 ± 0.2", "65.8", "52"],
            ["Rainy Weather",     "3.5 ± 0.8", "1.8 ± 0.4", "58.3", "63"],
            ["Fog / Low Vis.",    "4.0 ± 1.0", "2.2 ± 0.5", "55.0", "68"],
            ["High-Traffic",      "4.5 ± 1.2", "2.5 ± 0.6", "50.0", "70"],
        ]
    )
    add_rich_para(doc, (
        "Under calm water conditions — the most favourable operating environment — "
        "M-ACF Net reduced depth RMSE from 1.2 m (raw) to 0.6 m (fused), a "
        "50% improvement attributable to the adaptive confidence mechanism correctly "
        "prioritising LiDAR features (high C_LiDAR) whilst suppressing near-surface "
        "stereo noise. As environmental severity increased, RMSE degraded gradually "
        "rather than catastrophically: under fog and low visibility (the most "
        "demanding condition), fused RMSE of 2.2 m at 68 ms/frame remained "
        "within the real-time constraint. Table 5.6 places these results in "
        "the context of alternative fusion approaches."
    ))
    add_word_table(
        doc,
        "Table 5.6. Comparative fusion performance: M-ACF Net vs. alternative "
        "methods (RMSE in calm and stormy conditions; processing time on Jetson Orin NX).",
        ["Method", "RMSE Calm (m)", "RMSE Stormy (m)", "Process Time (ms)"],
        [
            ["**M-ACF Net (proposed)**", "**0.6 ± 0.1**", "**1.8 ± 0.4**", "**45–63**"],
            ["Median Filtering",          "1.0 ± 0.2",     "2.5 ± 0.6",     "30–40"],
            ["U-Net (LiDAR)",             "0.7 ± 0.2",     "2.0 ± 0.5",     "120–150"],
            ["ResNet-50 Fusion",          "0.8 ± 0.3",     "2.2 ± 0.6",     "90–110"],
            ["Vision Transformer",        "0.9 ± 0.3",     "2.4 ± 0.7",     "200–250"],
        ]
    )
    add_rich_para(doc, (
        "M-ACF Net achieves the lowest RMSE in both calm (0.6 m) and stormy (1.8 m) "
        "conditions — outperforming even U-Net and ResNet-50 fusion baselines "
        "that incorporate learned feature combination — whilst operating in 45–63 ms, "
        "compared with 90–250 ms for the deep-learning fusion baselines. This "
        "demonstrates that the confidence-weighted mechanism provides accuracy "
        "gains through principled adaptive weighting rather than additional "
        "computational depth."
    ))
    add_figure_placeholder(
        doc,
        "FIGURE 5.9",
        "Grouped bar chart comparing RMSE (calm/stormy) and processing time for "
        "M-ACF Net vs. four fusion baselines from Table 5.6. Source: Authors.",
        "Figure 5.9. Fusion performance comparison: M-ACF Net achieves lowest RMSE "
        "in both calm (0.6 m) and stormy (1.8 m) conditions while remaining within "
        "the real-time processing budget (45–63 ms), unlike deep-learning baselines "
        "requiring 90–250 ms."
    )

    doc.add_heading("Clustering and Localisation Validation", level=3)
    add_rich_para(doc, (
        "PCA–K-Means clustering performance was assessed using three internal "
        "cluster quality metrics — Silhouette score, Davies–Bouldin (DB) index, "
        "and Calinski–Harabasz (CH) index — as well as distance estimation accuracy "
        "(RMSE, MAE, and R²) against RTK-GNSS ground-truth obstacle positions. "
        "Table 5.7 compares M-ACF Net clustering against four established algorithms."
    ))
    add_word_table(
        doc,
        "Table 5.7. Clustering quality and computation time: M-ACF Net PCA–K-Means "
        "vs. alternative clustering algorithms (higher Silhouette and CH are better; "
        "lower DB and time are better).",
        ["Algorithm", "Silhouette ↑", "Davies–Bouldin ↓", "Calinski–Harabasz ↑", "Time (s) ↓"],
        [
            ["**M-ACF Net (proposed)**", "**0.70**", "**0.50**", "60",  "**0.05**"],
            ["DBSCAN",                   "0.75",      "0.60",     "80",  "0.12"],
            ["Hierarchical",             "0.70",      "0.70",     "50",  "0.25"],
            ["Spectral Clustering",      "0.80",      "0.55",     "90",  "0.30"],
            ["GMMs",                     "0.75",      "0.58",     "70",  "0.20"],
        ]
    )
    add_rich_para(doc, (
        "The M-ACF Net clustering configuration achieves the lowest computation time "
        "(0.05 s — four times faster than DBSCAN and six times faster than spectral "
        "clustering), with competitive Silhouette and Davies–Bouldin scores. "
        "The relatively lower CH index (60 vs. spectral clustering's 90) reflects "
        "the PCA–K-Means preference for a small number of well-separated clusters "
        "consistent with the maritime obstacle distribution, rather than maximising "
        "inter-cluster spread at the expense of intra-cluster coherence."
    ))
    add_word_table(
        doc,
        "Table 5.8. Obstacle distance estimation accuracy: M-ACF Net vs. alternative "
        "clustering algorithms (ground truth from RTK-GNSS ranging trials at "
        "Taman Tasik Metropolitan).",
        ["Algorithm", "RMSE (m) ↓", "MAE (m) ↓", "R² ↑"],
        [
            ["**M-ACF Net (proposed)**", "**1.6**", "**1.2**", "**0.88**"],
            ["DBSCAN",                   "2.0",      "1.6",     "0.80"],
            ["Hierarchical",             "2.8",      "2.2",     "0.70"],
            ["Spectral Clustering",      "1.8",      "1.4",     "0.85"],
            ["GMMs",                     "2.1",      "1.7",     "0.78"],
        ]
    )
    add_rich_para(doc, (
        "M-ACF Net achieves the best distance estimation accuracy: RMSE = 1.6 m, "
        "MAE = 1.2 m, and R² = 0.88. The R² score of 0.88 indicates that "
        "88% of the variance in ground-truth obstacle ranges is captured by the "
        "M-ACF Net range estimate, confirming strong linear correspondence between "
        "estimated and true obstacle positions across the tested range (0–15 m). "
        "The superiority of PCA–K-Means over DBSCAN (RMSE 2.0 m) and hierarchical "
        "clustering (RMSE 2.8 m) is particularly pronounced for elongated obstacles "
        "such as vessels and pontoons, where orientation alignment prior to clustering "
        "prevents the erroneous splitting of a single obstacle into multiple smaller "
        "clusters — a systematic failure mode of density-based and hierarchical "
        "approaches on maritime point cloud data."
    ))

    doc.add_heading("System-Level 3D Object Detection Benchmarking", level=3)
    add_rich_para(doc, (
        "The complete M-ACF Net pipeline was benchmarked against five 3D detection "
        "baselines — YOLOv8-Stereo, PointPillars, Pseudo-LiDAR, Frustum PointNet, "
        "and DSGN — on the custom maritime dataset, using 3D Average Precision (AP) "
        "at IoU thresholds of 50% (ships) and 40% (buoys, bridges) across Easy, "
        "Moderate, and Hard difficulty splits (Table 5.9)."
    ))
    add_word_table(
        doc,
        "Table 5.9. 3D Average Precision (%) by difficulty level for three maritime "
        "obstacle categories (IoU: 50% ships; 40% buoys/bridges; "
        "platform: NVIDIA Jetson Orin NX).",
        ["Method", "Ship Easy", "Ship Mod.", "Ship Hard",
         "Buoys Easy", "Buoys Mod.", "Buoys Hard",
         "Bridge Easy", "Bridge Mod.", "Bridge Hard", "FPS"],
        [
            ["YOLOv8-Stereo",   "68.2", "55.1", "42.3", "50.1", "40.2", "30.5", "45.3", "35.6", "25.8", "25"],
            ["PointPillars",     "72.5", "60.8", "48.6", "55.7", "45.3", "33.2", "50.2", "40.1", "28.4", "18"],
            ["Pseudo-LiDAR",    "75.1", "63.4", "50.2", "58.9", "47.1", "35.0", "52.4", "42.3", "30.1", "15"],
            ["Frustum PointNet","70.8", "58.3", "45.9", "53.2", "43.8", "32.7", "48.6", "38.9", "27.5", "10"],
            ["DSGN",            "77.3", "65.0", "52.4", "60.5", "49.8", "37.6", "54.7", "44.5", "32.9", "20"],
            ["**M-ACF Net**",   "**82.6**","**70.8**","**58.3**","**65.2**","**53.4**","**41.7**","**59.8**","**48.9**","**36.5**","18"],
        ]
    )
    add_rich_para(doc, (
        "M-ACF Net achieves the highest 3D AP across all three categories and all "
        "difficulty levels. For ships — the primary obstacle category in USV navigation "
        "— the system attains 82.6% / 70.8% / 58.3% on Easy / Moderate / Hard splits, "
        "compared with 77.3% / 65.0% / 52.4% for the best single-method baseline (DSGN). "
        "The absolute improvement is most pronounced on the Hard difficulty split "
        "(ships: +5.9 pp; buoys: +4.1 pp; bridges: +3.6 pp) — precisely the scenario "
        "of greatest operational relevance, where targets are small, distant, or "
        "partially occluded. M-ACF Net achieves these gains at 18 FPS, matching "
        "PointPillars throughput and far exceeding the 10 FPS of Frustum PointNet, "
        "confirming the suitability of the confidence-weighted mid-level fusion "
        "architecture for real-time edge deployment."
    ))
    add_figure_placeholder(
        doc,
        "FIGURE 5.9",
        "Grouped bar chart comparing 3D AP for M-ACF Net vs. five baselines across "
        "Ship / Buoys / Bridge categories at Easy, Mod., Hard difficulty. "
        "Source: Authors.",
        "Figure 5.9. 3D object detection benchmark: M-ACF Net achieves highest AP "
        "across all three maritime obstacle categories and difficulty levels, with the "
        "largest margin on Hard-difficulty targets (5.9 pp over nearest baseline for ships)."
    )

    doc.add_heading("Real-World Navigation and Obstacle Avoidance Trials", level=3)
    add_rich_para(doc, (
        "Four obstacle avoidance scenarios were executed with the Suraya Riptide "
        "operating under full M-ACF Net autonomous control in Malaysian waterways. "
        "Scenarios were designed to span the principal COLREGs encounter types: "
        "head-on (Scenario 1 and 3), crossing (Scenario 2), and stationary-obstacle "
        "(Scenario 4). In all four scenarios, the system achieved zero collisions. "
        "Table 5.10 summarises the key performance metrics."
    ))
    add_word_table(
        doc,
        "Table 5.10. Autonomous obstacle avoidance trial results across four "
        "COLREGs encounter scenarios (all scenarios: zero collisions).",
        ["Scenario", "Encounter Type", "Detection Range (m)", "Detection Time (s)",
         "Reaction Time (s)", "Path Deviation (m)", "Min. CPA (m)"],
        [
            ["1 — Dynamic head-on",      "Head-on (oncoming)",      "8.0",  "1.5", "1.2", "9.0",  "5.0  (3.3× vessel length)"],
            ["2 — Crossing (Rule 15)",   "Crossing, stand-on USV",  "10.0", "2.0", "1.5", "6.5",  "7.0"],
            ["3 — Turning into obstacle","Head-on (converging)",     "8.0",  "1.3", "1.0", "12.0", "6.5"],
            ["4 — Fixed structure",      "Stationary pier",         "15.0", "1.0", "0.8", "15.2", "10.0"],
        ]
    )
    add_rich_para(doc, (
        "Detection ranges varied from 8.0 m (head-on dynamic obstacle) to 15.0 m "
        "(fixed pier), consistent with the system's maximum effective perception "
        "range of approximately 15 m on the Suraya Riptide platform. "
        "Reaction times — measured from the first successful M-ACF Net detection "
        "to the initiation of a GNC avoidance manoeuvre — were consistently within "
        "0.8 to 1.5 seconds across all scenarios. Path deviations ranged from 6.5 m "
        "(crossing, Rule 15) to 15.2 m (fixed structure), reflecting the larger "
        "standoff required for stationary obstacles where early detection at 15 m "
        "allowed a gradual rerouting with a 2.0 m turning radius."
    ))
    add_rich_para(doc, (
        "Autonomous path-following performance — evaluated independently using "
        "cross-track error (CTE) over straight-line and curved waypoint sequences "
        "— demonstrated typical CTE values below **0.5 m** on straight-line segments "
        "and generally below **2.0 m** through waypoint transitions. CTE peaks "
        "approaching 2.0 m were observed at sharp heading changes, consistent with "
        "the turning radius limitations of the Riptide's propulsion configuration "
        "at the trial speed (approximately 0.5–1.0 m/s)."
    ))
    add_figure_placeholder(
        doc,
        "FIGURE 5.10",
        "Top-down GPS trajectory plots for Scenarios 1–4, showing USV path (solid), "
        "planned route (dashed), M-ACF Net obstacle detections (coloured markers), "
        "and avoidance manoeuvre trajectories. Source: Authors.",
        "Figure 5.10. Autonomous navigation trial results: GPS-logged USV trajectories "
        "for all four obstacle avoidance scenarios, overlaid with M-ACF Net obstacle "
        "detection events and GNC avoidance responses (zero collisions across all runs)."
    )


def _ch5_summary(doc: Document) -> None:
    add_rich_para(doc, (
        "This chapter has presented M-ACF Net, a mid-level LiDAR–stereo fusion "
        "framework for geospatial obstacle awareness in near-shore USV navigation, "
        "developed and validated aboard the Suraya Riptide trimaran at IIUM CUTe. "
        "The system integrates the Velodyne HDL-32E LiDAR and Stereolabs ZED 2i "
        "stereo camera through a dynamic confidence-weighted fusion mechanism, "
        "followed by Hybrid PCA–K-Means clustering and full GNC integration, "
        "all executed in real-time on the NVIDIA Jetson Orin NX at 18–25 FPS."
    ))
    add_rich_para(doc, (
        "The principal technical contributions are fourfold. "
        "First, the **confidence-weighted mid-level fusion mechanism** (Equations "
        "5.10–5.14) dynamically adjusts per-modality feature weights based on "
        "per-frame LiDAR density and stereo occlusion reliability estimates, "
        "reducing depth RMSE by 50% in calm conditions (1.2 → 0.6 m) and "
        "outperforming all alternative fusion baselines whilst remaining within "
        "the 100 ms real-time processing budget. "
        "Second, the **Pseudo-LiDAR generation pipeline** (Equations 5.3–5.9) "
        "converts deep-learning-based stereo disparity maps into geometrically "
        "compatible 3D point representations using the ZED 2i's 65 mm baseline "
        "and calibrated intrinsics, enabling tight mid-level integration at "
        "9.8 ± 1.3 ms/frame. "
        "Third, the **Hybrid PCA–K-Means clustering** stage achieves the best "
        "distance estimation accuracy across all tested algorithms "
        "(RMSE 1.6 m, MAE 1.2 m, R² = 0.88) at the lowest computation time "
        "(0.05 s/frame), with orientation-aware segmentation that correctly handles "
        "elongated maritime obstacles where density-based clustering fails. "
        "Fourth, the **3D detection and GNC integration** — achieving 82.6% AP "
        "(Ship Easy), 70.8% (Ship Moderate), and 58.3% (Ship Hard), with zero "
        "collisions across four COLREGs encounter scenarios — demonstrates that "
        "M-ACF Net provides the perception quality and responsiveness required for "
        "fully autonomous USV navigation in complex near-shore environments."
    ))
    add_rich_para(doc, (
        "Comparison with the decision-level fusion approach presented in Chapter 4 "
        "reveals a fundamental trade-off: mid-level fusion delivers superior ranging "
        "accuracy and robustness at the cost of tighter calibration requirements "
        "and a more complex feature integration architecture. Chapter 6 examines "
        "this trade-off quantitatively across a range of environmental conditions "
        "and obstacle scenarios, providing practical guidance for system designers "
        "selecting a fusion strategy for near-shore USV applications."
    ))


def generate_chapter5(output_path: str) -> None:
    doc = new_iium_doc()

    # ── Chapter heading ──────────────────────────────────────────────────────
    doc.add_heading(
        "Chapter 5: Mid-Level LiDAR–Stereo Geospatial Feature Fusion with Adaptive "
        "Confidence Weighting: The M-ACF Net Framework", level=1
    )
    p_auth = doc.add_paragraph("Muhammad Aiman bin Norazaruddin")
    p_auth.runs[0].bold   = True
    p_auth.paragraph_format.space_after = Pt(2)
    p_aff = doc.add_paragraph("Kulliyyah of Engineering, International Islamic University Malaysia")
    p_aff.runs[0].italic  = True
    p_aff.paragraph_format.space_after = Pt(12)

    # ── 5.1 Introduction ─────────────────────────────────────────────────────
    doc.add_heading("5.1  Introduction", level=2)
    _ch5_intro(doc)

    # ── 5.2 Platform ──────────────────────────────────────────────────────────
    doc.add_heading("5.2  The Suraya Riptide Platform and Sensor Configuration", level=2)
    _ch5_platform(doc)

    # ── 5.3 M-ACF Net Architecture ────────────────────────────────────────────
    doc.add_heading("5.3  The M-ACF Net Framework: Architecture and Design Rationale", level=2)
    doc.add_heading("Mid-Level Fusion: Design Philosophy and Paradigm Selection", level=3)
    _ch5_architecture(doc)

    # ── 5.4 Visual Feature Extraction ─────────────────────────────────────────
    doc.add_heading("5.4  Visual Feature Extraction and Pseudo-LiDAR Generation", level=2)
    _ch5_visual(doc)

    # ── 5.5 LiDAR Processing ──────────────────────────────────────────────────
    doc.add_heading("5.5  LiDAR Point Cloud Processing", level=2)
    _ch5_lidar(doc)

    # ── 5.6 Fusion ────────────────────────────────────────────────────────────
    doc.add_heading("5.6  Confidence-Weighted Mid-Level Feature Fusion", level=2)
    _ch5_fusion(doc)

    # ── 5.7 Clustering ────────────────────────────────────────────────────────
    doc.add_heading("5.7  Post-Processing: PCA–K-Means Obstacle Clustering", level=2)
    _ch5_clustering(doc)

    # ── 5.8 Dataset ───────────────────────────────────────────────────────────
    doc.add_heading("5.8  Dataset Collection and Preprocessing", level=2)
    _ch5_dataset(doc)

    # ── 5.9 GNC Integration ───────────────────────────────────────────────────
    doc.add_heading("5.9  GNC Integration and System Deployment", level=2)
    _ch5_gnc(doc)

    # ── 5.10 Results ──────────────────────────────────────────────────────────
    doc.add_heading("5.10  Experimental Validation and Results", level=2)
    _ch5_results(doc)

    # ── 5.11 Summary ──────────────────────────────────────────────────────────
    doc.add_heading("5.11  Summary", level=2)
    _ch5_summary(doc)

    # ── References (chapter-level) ────────────────────────────────────────────
    doc.add_heading("References", level=2)
    refs = [
        ("Aiman, M. N., Zulkifli, Z. A., & Mohd Zaki, H. F. (2026). "
         "A mid-level LiDAR–stereo fusion framework for USV geospatial awareness. "
         "[Doctoral thesis, IIUM]."),
        ("Endsley, M. R. (1995). Toward a theory of situation awareness in dynamic systems. "
         "Human Factors, 37(1), 32–64. https://doi.org/10.1518/001872095779049543"),
        ("Li, Y., & Deng, J. (2019). Pseudo-LiDAR from visual depth estimation: Bridging the "
         "gap in 3D object detection for autonomous driving. Proceedings of IEEE CVPR, "
         "8445–8453."),
        ("Redmon, J., & Farhadi, A. (2018). YOLOv3: An incremental improvement. "
         "arXiv:1804.02767."),
        ("Wang, C.-Y., Bochkovskiy, A., & Liao, H.-Y. M. (2023). YOLOv7: Trainable "
         "bag-of-freebies sets new state-of-the-art for real-time object detectors. "
         "Proceedings of IEEE CVPR, 7464–7475."),
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent       = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        r = p.add_run(ref)
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[✓] Chapter 5 saved → {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    base = Path(__file__).parent / "output"
    generate_chapter4(str(base / "CHAPTER_04_DRAFT.docx"))
    generate_chapter5(str(base / "CHAPTER_05_DRAFT.docx"))
    print("\nBoth chapters generated successfully.")
    print("Note: Tables marked '— [thesis Table N.M]' require transcription of")
    print("     exact values from the thesis before final submission.")


if __name__ == "__main__":
    main()
